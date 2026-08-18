"""Render CV data and cover letters to professional, ATS-friendly PDF and DOCX.

Design goals (blueprint sections 32 & 33):
- Single column, real selectable text, standard section headings, no tables /
  text boxes / images — the layout ATS parsers read most reliably.
- Polished and HR-attractive: a clean branded header, clear section rules,
  achievement bullet points, and consistent typography.
PDF via ReportLab, DOCX via python-docx.
"""
from __future__ import annotations

import io
import re
from datetime import date
from xml.sax.saxutils import escape

from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, HRFlowable, ListFlowable, ListItem,
)

# Brand palette (kept subtle so the document still reads as a serious CV).
NAVY = colors.HexColor("#0b2447")
TEAL = colors.HexColor("#0f766e")
GOLD = colors.HexColor("#f5b301")
GREY = colors.HexColor("#5b6675")


def _contact_line(cv: dict) -> str:
    bits = [cv.get("email"), cv.get("phone")]
    loc = ", ".join(x for x in [cv.get("city"), cv.get("country")] if x)
    if loc:
        bits.append(loc)
    bits += [cv.get("linkedin_url"), cv.get("github_url"), cv.get("portfolio_url")]
    return "  |  ".join(b for b in bits if b)


def _exp_dates(e: dict) -> str:
    start = e.get("start_date") or ""
    end = "Present" if e.get("is_current") else (e.get("end_date") or "")
    return " – ".join(x for x in [str(start), str(end)] if x)


def _bullets(text) -> list[str]:
    """Split a free-text responsibilities/achievements field into clean bullets.

    Splits on newlines, semicolons, and bullet characters, and on sentence
    boundaries only when the text is one long run. Truthful: it reorganises the
    candidate's own words, never adds content.
    """
    if not text:
        return []
    raw = str(text).strip()
    parts = re.split(r"[\n;•·]|(?:(?<=[a-z0-9\)])\.\s+(?=[A-Z]))", raw)
    out = []
    for p in parts:
        s = (p or "").strip(" \t-–—.")
        if len(s) >= 2:
            out.append(s[0].upper() + s[1:])
    return out


# ---- PDF (ReportLab) --------------------------------------------------------

def render_cv_pdf(cv: dict) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=1.5 * cm, bottomMargin=1.5 * cm,
                            leftMargin=1.9 * cm, rightMargin=1.9 * cm,
                            title=(cv.get("full_name") or "Curriculum Vitae"))
    base = getSampleStyleSheet()
    name_style = ParagraphStyle("Name", parent=base["Title"], fontName="Helvetica-Bold",
                                fontSize=22, textColor=NAVY, spaceAfter=1, alignment=TA_LEFT, leading=24)
    role_style = ParagraphStyle("Role", parent=base["Normal"], fontName="Helvetica-Bold",
                                fontSize=11.5, textColor=TEAL, spaceAfter=3)
    contact_style = ParagraphStyle("Contact", parent=base["Normal"], fontSize=9, textColor=GREY, leading=13)
    section_style = ParagraphStyle("Section", parent=base["Heading2"], fontName="Helvetica-Bold",
                                   fontSize=11, textColor=NAVY, spaceBefore=12, spaceAfter=3, leading=13)
    body = ParagraphStyle("Body", parent=base["Normal"], fontSize=10, leading=14.5)
    job_style = ParagraphStyle("Job", parent=body, fontName="Helvetica-Bold", fontSize=10.5, spaceBefore=6)
    meta_style = ParagraphStyle("Meta", parent=base["Normal"], fontSize=9, textColor=GREY, spaceAfter=2)
    bullet_style = ParagraphStyle("Bullet", parent=body, fontSize=10, leading=14)

    def section(title: str):
        return [Paragraph(title.upper(), section_style),
                HRFlowable(width="100%", thickness=1, color=GOLD, spaceBefore=1, spaceAfter=5)]

    def bullet_list(items: list[str]):
        return ListFlowable(
            [ListItem(Paragraph(escape(i), bullet_style), leftIndent=10, value=None) for i in items],
            bulletType="bullet", bulletChar="•", bulletColor=TEAL, bulletFontSize=9,
            leftIndent=12, spaceBefore=1, spaceAfter=1,
        )

    story = [Paragraph(escape(cv.get("full_name") or "Curriculum Vitae"), name_style)]
    role = cv.get("target_vacancy_title") or cv.get("current_occupation")
    if role and str(role).strip().lower() != "the role":
        story.append(Paragraph(escape(str(role)), role_style))
    contact = _contact_line(cv)
    if contact:
        story.append(Paragraph(escape(contact), contact_style))
    story.append(HRFlowable(width="100%", thickness=2, color=NAVY, spaceBefore=6, spaceAfter=2))

    if cv.get("summary"):
        story += section("Professional Summary")
        story.append(Paragraph(escape(cv["summary"]), body))

    if cv.get("skills"):
        story += section("Core Skills")
        story.append(Paragraph(escape("  •  ".join(cv["skills"])), body))

    if cv.get("experience"):
        story += section("Work Experience")
        for e in cv["experience"]:
            header = " — ".join(x for x in [e.get("position"), e.get("employer")] if x)
            if header:
                story.append(Paragraph(escape(header), job_style))
            meta = "  ·  ".join(x for x in [_exp_dates(e), e.get("industry")] if x)
            if meta:
                story.append(Paragraph(escape(meta), meta_style))
            resp = _bullets(e.get("responsibilities"))
            if resp:
                story.append(bullet_list(resp))
            achv = _bullets(e.get("achievements"))
            if achv:
                story.append(Paragraph("<b>Key achievements</b>", meta_style))
                story.append(bullet_list(achv))

    if cv.get("education"):
        story += section("Education")
        for ed in cv["education"]:
            line = " — ".join(x for x in [ed.get("qualification"), ed.get("institution")] if x)
            if ed.get("completion_date"):
                line += f"  ({ed['completion_date']})"
            story.append(Paragraph(escape(line or ed.get("institution", "")), body))

    if cv.get("certifications"):
        story += section("Certifications")
        for c in cv["certifications"]:
            line = " — ".join(x for x in [c.get("name"), c.get("issuing_organization")] if x)
            story.append(Paragraph(escape(line), body))

    extras = []
    if cv.get("languages"):
        extras.append("Languages: " + ", ".join(cv["languages"]))
    if cv.get("drivers_licence"):
        extras.append("Driver's licence: " + str(cv["drivers_licence"]))
    if extras:
        story += section("Additional")
        for x in extras:
            story.append(Paragraph(escape(x), body))

    doc.build(story)
    return buf.getvalue()


# ---- DOCX (python-docx) -----------------------------------------------------

def render_cv_docx(cv: dict) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor

    NAVY_D = RGBColor(0x0B, 0x24, 0x47)
    TEAL_D = RGBColor(0x0F, 0x76, 0x6E)
    GREY_D = RGBColor(0x5B, 0x66, 0x75)

    d = Document()
    normal = d.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    # Header
    p = d.add_paragraph()
    r = p.add_run(cv.get("full_name") or "Curriculum Vitae")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = NAVY_D
    role = cv.get("target_vacancy_title") or cv.get("current_occupation")
    if role and str(role).strip().lower() != "the role":
        rp = d.add_paragraph()
        rr = rp.add_run(str(role))
        rr.bold = True
        rr.font.size = Pt(12)
        rr.font.color.rgb = TEAL_D
    contact = _contact_line(cv)
    if contact:
        cp = d.add_paragraph()
        cr = cp.add_run(contact)
        cr.font.size = Pt(9)
        cr.font.color.rgb = GREY_D

    def section(title: str):
        h = d.add_paragraph()
        hr = h.add_run(title.upper())
        hr.bold = True
        hr.font.size = Pt(11)
        hr.font.color.rgb = NAVY_D
        _bottom_border(d.add_paragraph())

    def bullets(items):
        for i in items:
            d.add_paragraph(i, style="List Bullet")

    if cv.get("summary"):
        section("Professional Summary")
        d.add_paragraph(cv["summary"])
    if cv.get("skills"):
        section("Core Skills")
        d.add_paragraph("  •  ".join(cv["skills"]))
    if cv.get("experience"):
        section("Work Experience")
        for e in cv["experience"]:
            header = " — ".join(x for x in [e.get("position"), e.get("employer")] if x)
            hp = d.add_paragraph()
            hr = hp.add_run(header)
            hr.bold = True
            meta = "  ·  ".join(x for x in [_exp_dates(e), e.get("industry")] if x)
            if meta:
                mp = d.add_paragraph()
                mr = mp.add_run(meta)
                mr.font.size = Pt(9)
                mr.font.color.rgb = GREY_D
            bullets(_bullets(e.get("responsibilities")))
            achv = _bullets(e.get("achievements"))
            if achv:
                ap = d.add_paragraph()
                ar = ap.add_run("Key achievements")
                ar.bold = True
                ar.font.size = Pt(9)
                bullets(achv)
    if cv.get("education"):
        section("Education")
        for ed in cv["education"]:
            line = " — ".join(x for x in [ed.get("qualification"), ed.get("institution")] if x)
            if ed.get("completion_date"):
                line += f"  ({ed['completion_date']})"
            d.add_paragraph(line or ed.get("institution", ""))
    if cv.get("certifications"):
        section("Certifications")
        for c in cv["certifications"]:
            d.add_paragraph(" — ".join(x for x in [c.get("name"), c.get("issuing_organization")] if x))
    if cv.get("languages") or cv.get("drivers_licence"):
        section("Additional")
        if cv.get("languages"):
            d.add_paragraph("Languages: " + ", ".join(cv["languages"]))
        if cv.get("drivers_licence"):
            d.add_paragraph("Driver's licence: " + str(cv["drivers_licence"]))

    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _bottom_border(paragraph):
    """Add a thin bottom border to a paragraph (used as a section rule)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "f5b301")
    pbdr.append(bottom)
    pPr.append(pbdr)


# ---- Cover letter -----------------------------------------------------------

def render_letter_pdf(text: str) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=2 * cm, bottomMargin=2 * cm,
                            leftMargin=2.2 * cm, rightMargin=2.2 * cm, title="Cover Letter")
    base = getSampleStyleSheet()
    body = ParagraphStyle("LetterBody", parent=base["Normal"], fontSize=11, leading=16, spaceAfter=8)
    story = []
    for para in text.split("\n\n"):
        story.append(Paragraph(escape(para).replace("\n", "<br/>"), body))
    doc.build(story)
    return buf.getvalue()


def render_letter_docx(text: str) -> bytes:
    from docx import Document
    from docx.shared import Pt
    d = Document()
    d.styles["Normal"].font.name = "Calibri"
    d.styles["Normal"].font.size = Pt(11)
    for para in text.split("\n\n"):
        d.add_paragraph(para)
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()
