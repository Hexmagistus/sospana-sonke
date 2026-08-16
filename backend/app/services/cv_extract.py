"""CV text extraction (blueprint section 4).

Turns an uploaded PDF/DOCX/TXT into plain text. AI structuring runs on this text.
Extraction failures are raised as CVExtractionError so the API can mark the CV as
failed rather than silently losing it (blueprint section 35: never fail silently).
"""
from __future__ import annotations

import io


class CVExtractionError(Exception):
    pass


def _extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(data))
        parts = [(page.extract_text() or "") for page in reader.pages]
        text = "\n".join(parts).strip()
    except Exception as exc:
        raise CVExtractionError(f"Could not read PDF: {exc}") from exc
    if not text:
        raise CVExtractionError(
            "No text found in PDF. It may be a scanned image; OCR is a later-phase feature."
        )
    return text


def _extract_docx(data: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(data))
        lines = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                lines.append(" ".join(cell.text for cell in row.cells))
        text = "\n".join(line for line in lines if line is not None).strip()
    except Exception as exc:
        raise CVExtractionError(f"Could not read DOCX: {exc}") from exc
    if not text:
        raise CVExtractionError("No text found in DOCX.")
    return text


def _extract_txt(data: bytes) -> str:
    for encoding in ("utf-8", "latin-1"):
        try:
            text = data.decode(encoding).strip()
            if text:
                return text
        except UnicodeDecodeError:
            continue
    raise CVExtractionError("Could not decode text file.")


def extract_text(data: bytes, extension: str) -> str:
    if extension == "pdf":
        return _extract_pdf(data)
    if extension == "docx":
        return _extract_docx(data)
    if extension == "txt":
        return _extract_txt(data)
    raise CVExtractionError(f"Unsupported extension '.{extension}'.")
